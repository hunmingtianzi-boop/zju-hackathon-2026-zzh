import tempfile
import unittest
from pathlib import Path

from multi_source_loader import DocumentLoader, markdown_table


class DocumentLoaderTest(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.root = Path(self.temp_dir.name)
        self.loader = DocumentLoader()

    def tearDown(self):
        self.temp_dir.cleanup()

    def test_markdown_sections_are_split_by_headings(self):
        path = self.root / "sample.md"
        path.write_text("# Book\n\nIntro\n\n## Chapter A\n\nAlpha\n\n## Chapter B\n\nBeta", encoding="utf-8")

        sections = self.loader.load(path)

        self.assertEqual([section.title for section in sections], ["Book", "Chapter A", "Chapter B"])
        self.assertEqual(sections[1].source_type, "md")
        self.assertIn("Alpha", sections[1].markdown)

    def test_word_document_is_loaded_as_markdown_sections(self):
        from docx import Document

        path = self.root / "notes.docx"
        document = Document()
        document.add_heading("Clinical Notes", level=1)
        document.add_paragraph("A paragraph from Word.")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Term"
        table.cell(0, 1).text = "Meaning"
        table.cell(1, 0).text = "HR"
        table.cell(1, 1).text = "Heart rate"
        document.save(path)

        sections = self.loader.load(path)

        self.assertEqual(len(sections), 1)
        self.assertEqual(sections[0].title, "Clinical Notes")
        self.assertIn("| Term | Meaning |", sections[0].markdown)

    def test_excel_workbook_creates_one_section_per_sheet(self):
        from openpyxl import Workbook

        path = self.root / "tables.xlsx"
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Vitals"
        sheet.append(["Metric", "Value"])
        sheet.append(["BP", "120/80"])
        second = workbook.create_sheet("Empty")
        second.append([None, None])
        workbook.save(path)

        sections = self.loader.load(path)

        self.assertEqual(len(sections), 1)
        self.assertEqual(sections[0].title, "Vitals")
        self.assertEqual(sections[0].source_type, "xlsx")
        self.assertIn("120/80", sections[0].markdown)

    def test_pdf_pages_include_page_metadata(self):
        import fitz

        path = self.root / "mini.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Chapter 1 Overview", fontsize=18)
        page.insert_text((72, 110), "PDF textbook text", fontsize=11)
        doc.save(path)
        doc.close()

        sections = self.loader.load(path)

        self.assertEqual(len(sections), 1)
        self.assertEqual(sections[0].title, "Chapter 1 Overview")
        self.assertEqual(sections[0].page_start, 1)
        self.assertIn("PDF textbook text", sections[0].markdown)
        self.assertIn("section_id", sections[0].metadata)

    def test_manifest_summarizes_loaded_sources(self):
        path = self.root / "sample.md"
        path.write_text("# Book\n\nIntro\n\n## Chapter A\n\nAlpha", encoding="utf-8")

        sections = self.loader.load(path)
        manifest = self.loader.build_manifest(sections)

        self.assertEqual(manifest["total_books"], 1)
        self.assertEqual(manifest["books"][0]["sections"], 2)
        self.assertEqual(manifest["books"][0]["source_type"], "md")

    def test_load_many_discovers_supported_files(self):
        (self.root / "a.md").write_text("# A\n\nAlpha", encoding="utf-8")
        (self.root / "ignore.txt").write_text("ignored", encoding="utf-8")

        sections = self.loader.load_many(self.root)

        self.assertEqual(len(sections), 1)
        self.assertEqual(sections[0].book_id, "a")

    def test_markdown_table_normalizes_ragged_rows(self):
        table = markdown_table([["A", "B"], ["1"]])

        self.assertIn("| A | B |", table)
        self.assertIn("| 1 |  |", table)


if __name__ == "__main__":
    unittest.main()
